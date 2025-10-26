# reports_bp.py
from flask import Blueprint, render_template, session
from datetime import datetime
from core.db_utils import get_sql_connection
from core.decorators import require_role
from flask import Blueprint, send_file
from core.decorators import require_role
import io
from docx import Document
from openpyxl import Workbook

reports_bp = Blueprint("reports_bp", __name__)


# ======================================================
# 📊 1. BÁO CÁO THỐNG KÊ (ADMIN)
# ======================================================
@reports_bp.route('/reports')
@require_role("admin")
def reports_view():
    conn = get_sql_connection()
    cursor = conn.cursor()

    # --- 1. Danh sách nhân viên ---
    cursor.execute("""
        SELECT NV.*, PB.TenPB,
               CASE NV.TrangThai 
                   WHEN 1 THEN N'Đang hoạt động'
                   ELSE N'Ngừng hoạt động'
               END AS TrangThaiText
        FROM NhanVien NV
        LEFT JOIN PhongBan PB ON NV.MaPB = PB.MaPB
    """)
    columns = [col[0] for col in cursor.description]
    employees = [dict(zip(columns, row)) for row in cursor.fetchall()]

    # --- 2. Danh sách phòng ban ---
    cursor.execute("""
        SELECT PB.MaPB, PB.TenPB,
               COUNT(NV.MaNV) AS SoNhanVien,
               CASE 
                   WHEN LTRIM(RTRIM(LOWER(PB.TrangThai))) IN 
                        (N'đang hoạt động', N'active', N'1') 
                        THEN N'Đang hoạt động'
                   ELSE N'Ngừng hoạt động'
               END AS TrangThaiText
        FROM PhongBan PB
        LEFT JOIN NhanVien NV ON PB.MaPB = NV.MaPB
        GROUP BY PB.MaPB, PB.TenPB, PB.TrangThai
        ORDER BY PB.MaPB
    """)
    columns_pb = [col[0] for col in cursor.description]
    departments = [dict(zip(columns_pb, row)) for row in cursor.fetchall()]

    # --- 3. Thống kê ---
    total_employees = len(employees)
    total_departments = len(departments)

    cursor.execute("""
        SELECT 
            COUNT(DISTINCT MaNV) AS SoNhanVienCoLuong,
            SUM(TongTien) AS TongLuong,
            AVG(TongTien) AS LuongTrungBinh
        FROM Luong
        WHERE TrangThai = 1 AND DaXoa = 1
    """)
    salary_stats = cursor.fetchone()
    total_salary = salary_stats[1] or 0
    avg_salary = salary_stats[2] or 0

    cursor.execute("SELECT AVG(DATEDIFF(hour, GioVao, GioRa)) FROM ChamCong WHERE TrangThai = 1")
    avg_hours_per_week = cursor.fetchone()[0] or 0

    status_counts = {'Đang hoạt động': 0, 'Ngừng hoạt động': 0}
    for emp in employees:
        status_counts[emp['TrangThaiText']] += 1

    dept_counts, gender_counts = {}, {'Nam': 0, 'Nữ': 0}
    for emp in employees:
        dept_counts[emp.get('TenPB') or 'Chưa phân công'] = dept_counts.get(emp.get('TenPB') or 'Chưa phân công', 0) + 1
        if emp.get('GioiTinh') == 1:
            gender_counts['Nam'] += 1
        else:
            gender_counts['Nữ'] += 1

    cursor.execute("""
        SELECT CL.TenCa, COUNT(LLV.MaNV)
        FROM LichLamViec LLV
        JOIN CaLamViec CL ON LLV.MaCa = CL.MaCa
        GROUP BY CL.TenCa
    """)
    shift_counts = dict(cursor.fetchall())
    conn.close()

    return render_template(
        'reports.html',
        employees=employees,
        departments=departments,
        total_employees=total_employees,
        total_departments=total_departments,
        avg_hours_per_week=avg_hours_per_week,
        total_salary=total_salary,
        avg_salary=avg_salary,
        status_counts=status_counts,
        dept_counts=dept_counts,
        gender_counts=gender_counts,
        shift_counts=shift_counts
    )


# ======================================================
# 📊 2. BÁO CÁO HR
# ======================================================
@reports_bp.route('/hr/reports')
@require_role("hr")
def hr_reports_view():
    conn = get_sql_connection()
    cursor = conn.cursor()

    cursor.execute("""
        SELECT NV.MaNV, NV.HoTen, NV.Email, NV.GioiTinh, NV.TrangThai,
            PB.TenPB, NV.ChucVu
        FROM NhanVien NV
        LEFT JOIN PhongBan PB ON NV.MaPB = PB.MaPB
        WHERE NV.TrangThai IN (0, 1)
    """)
    columns = [col[0] for col in cursor.description]
    employees = [dict(zip(columns, row)) for row in cursor.fetchall()]

    cursor.execute("""
        SELECT PB.MaPB, PB.TenPB,
               COUNT(NV.MaNV) AS SoNhanVien,
               CASE 
                   WHEN PB.TrangThai = 1 THEN N'Đang hoạt động'
                   ELSE N'Ngừng hoạt động'
               END AS TrangThaiText
        FROM PhongBan PB
        LEFT JOIN NhanVien NV ON PB.MaPB = NV.MaPB
        GROUP BY PB.MaPB, PB.TenPB, PB.TrangThai
        ORDER BY PB.MaPB
    """)
    columns_pb = [col[0] for col in cursor.description]
    departments = [dict(zip(columns_pb, row)) for row in cursor.fetchall()]

    total_employees = len(employees)
    total_departments = len(departments)

    cursor.execute("""
        SELECT COUNT(*) AS TongCong,
               SUM(CASE WHEN TrangThai = 1 THEN 1 ELSE 0 END) AS DiLam
        FROM ChamCong
        WHERE DaXoa = 1
    """)
    row = cursor.fetchone()
    attendance_rate = round((row[1] or 0) / (row[0] or 1) * 100, 2)

    cursor.execute("""
        SELECT AVG(DATEDIFF(HOUR, GioVao, GioRa))
        FROM ChamCong
        WHERE GioVao IS NOT NULL AND GioRa IS NOT NULL
    """)
    avg_hours_per_week = round(cursor.fetchone()[0] or 0, 1)

    cursor.execute("SELECT COUNT(*) FROM CaLamViec WHERE TrangThai = 1")
    total_shifts = cursor.fetchone()[0] or 0

    cursor.execute("SELECT SUM(TongTien) FROM Luong WHERE TrangThai = 1 AND DaXoa = 1")
    total_salary = int(cursor.fetchone()[0] or 0)

    # Biểu đồ
    status_counts = {'Đang hoạt động': 0, 'Ngừng hoạt động': 0}
    gender_counts = {'Nam': 0, 'Nữ': 0}
    dept_counts = {}

    for emp in employees:
        if emp.get("TrangThai") == 1:
            status_counts['Đang hoạt động'] += 1
        else:
            status_counts['Ngừng hoạt động'] += 1

        if emp.get("GioiTinh") == 1:
            gender_counts['Nam'] += 1
        else:
            gender_counts['Nữ'] += 1

        dept_name = emp.get("TenPB") or "Chưa phân công"
        dept_counts[dept_name] = dept_counts.get(dept_name, 0) + 1

    cursor.execute("""
        SELECT CLV.TenCa, COUNT(LLV.MaNV)
        FROM LichLamViec LLV
        JOIN CaLamViec CLV ON LLV.MaCa = CLV.MaCa
        GROUP BY CLV.TenCa
    """)
    shift_counts = dict(cursor.fetchall())
    conn.close()

    return render_template(
        'hr_reports.html',
        employees=employees,
        departments=departments,
        total_employees=total_employees,
        total_departments=total_departments,
        attendance_rate=attendance_rate,
        avg_hours_per_week=avg_hours_per_week,
        total_shifts=total_shifts,
        total_salary=total_salary,
        status_counts=status_counts,
        gender_counts=gender_counts,
        dept_counts=dept_counts,
        shift_counts=shift_counts
    )
# ======================================================
# 📊 BÁO CÁO NHÂN SỰ & CHẤM CÔNG CHO QUẢN LÝ PHÒNG BAN
# ======================================================
@reports_bp.route("/qlpb/reports")
@require_role("quanlyphongban")
def qlpb_reports_view():
    conn = get_sql_connection()
    cursor = conn.cursor()
    username = session.get("username")

    # 🔹 1️⃣ Xác định mã phòng ban của quản lý
    cursor.execute("""
        SELECT NV.MaPB
        FROM NhanVien NV
        JOIN TaiKhoan TK ON NV.MaNV = TK.MaNV
        WHERE TK.TenDangNhap = ?
    """, (username,))
    row = cursor.fetchone()
    ma_pb_user = row[0] if row else None

    if not ma_pb_user:
        flash("❌ Không tìm thấy phòng ban của quản lý!", "danger")
        return redirect(url_for("qlpb_bp.qlpb_dashboard"))

    # 🔹 2️⃣ Lấy danh sách nhân viên thuộc phòng ban đó
    cursor.execute("""
        SELECT NV.MaNV, NV.HoTen, NV.Email, NV.GioiTinh, 
               PB.TenPB, NV.TrangThai,
               CASE NV.TrangThai 
                   WHEN 1 THEN N'Đang làm việc'
                   ELSE N'Ngừng làm việc'
               END AS TrangThaiText
        FROM NhanVien NV
        JOIN PhongBan PB ON NV.MaPB = PB.MaPB
        WHERE NV.MaPB = ?
    """, (ma_pb_user,))
    columns = [col[0] for col in cursor.description]
    employees = [dict(zip(columns, row)) for row in cursor.fetchall()]

    # 🔹 3️⃣ Tổng số nhân viên
    total_employees = len(employees)

    # 🔹 4️⃣ Thống kê chấm công trong tháng hiện tại
    cursor.execute("""
        SELECT 
            SUM(CASE WHEN LLV.TrangThai = 1 THEN 1 ELSE 0 END) AS SoCaDiLam,
            SUM(CASE WHEN LLV.TrangThai = 2 THEN 1 ELSE 0 END) AS SoCaVang,
            COUNT(*) AS TongCa
        FROM LichLamViec LLV
        JOIN NhanVien NV ON LLV.MaNV = NV.MaNV
        WHERE NV.MaPB = ? 
          AND MONTH(LLV.NgayLam) = MONTH(GETDATE())
          AND YEAR(LLV.NgayLam) = YEAR(GETDATE())
    """, (ma_pb_user,))
    stats = cursor.fetchone()
    so_ca_di_lam = stats[0] or 0
    so_ca_vang = stats[1] or 0
    tong_ca = stats[2] or 0

    # 🔹 5️⃣ Tính tỷ lệ chấm công
    attendance_rate = round((so_ca_di_lam / tong_ca * 100), 2) if tong_ca else 0

    # 🔹 6️⃣ Biểu đồ trạng thái làm việc
    status_counts = {
        "Đang làm việc": sum(1 for e in employees if e["TrangThai"] == 1),
        "Ngừng làm việc": sum(1 for e in employees if e["TrangThai"] == 0),
    }

    # 🔹 7️⃣ Biểu đồ giới tính
    gender_counts = {
        "Nam": sum(1 for e in employees if e["GioiTinh"] in (1, "Nam", "nam")),
        "Nữ": sum(1 for e in employees if e["GioiTinh"] in (0, "Nữ", "nu", "nữ")),
    }

    # 🔹 8️⃣ Biểu đồ phòng ban (chỉ 1 phòng nhưng vẫn hiển thị)
    dept_counts = {}
    for e in employees:
            pb = e.get("TenPB") or "Chưa phân công"
            dept_counts[pb] = dept_counts.get(pb, 0) + 1

        # 🔹 10️⃣ Biểu đồ số nhân viên theo ca làm
    cursor.execute("""
        SELECT CL.TenCa, COUNT(LLV.MaNV)
        FROM LichLamViec LLV
        JOIN CaLamViec CL ON LLV.MaCa = CL.MaCa
        JOIN NhanVien NV ON NV.MaNV = LLV.MaNV
        WHERE NV.MaPB = ?
        AND MONTH(LLV.NgayLam) = MONTH(GETDATE())
        AND YEAR(LLV.NgayLam) = YEAR(GETDATE())
        GROUP BY CL.TenCa
    """, (ma_pb_user,))
    shift_counts = dict(cursor.fetchall())

    conn.close()

    # 🔹 9️⃣ Trả về template
    return render_template(
        "qlpb_reports.html",
        employees=employees,
        total_employees=total_employees,
        so_ca_di_lam=so_ca_di_lam,
        so_ca_vang=so_ca_vang,
        attendance_rate=attendance_rate,
        status_counts=status_counts,
        gender_counts=gender_counts,
        dept_counts=dept_counts,
        shift_counts=shift_counts 
    )


# ============================================================
# 📊 XUẤT BÁO CÁO DẠNG WORD
# ============================================================
@reports_bp.route("/export_report/word")
@require_role("admin", "hr")
def export_report_word():
    # (Dữ liệu mẫu, bạn có thể thay bằng dữ liệu thật từ DB)
    total_employees = 6
    total_departments = 4
    attendance_rate = "96.5%"
    avg_hours = 4
    total_shifts = 1
    total_salary = 0

    # Tạo file Word
    doc = Document()
    doc.add_heading("BÁO CÁO TỔNG QUAN HỆ THỐNG", level=1)
    doc.add_paragraph(f"Tổng nhân viên: {total_employees}")
    doc.add_paragraph(f"Tổng phòng ban: {total_departments}")
    doc.add_paragraph(f"Tỉ lệ chấm công: {attendance_rate}")
    doc.add_paragraph(f"Giờ trung bình/Tuần: {avg_hours}")
    doc.add_paragraph(f"Tổng số ca: {total_shifts}")
    doc.add_paragraph(f"Tổng lương: {total_salary:,}")

    # Lưu ra bộ nhớ tạm
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="BaoCaoTongQuan.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ============================================================
# 📊 XUẤT BÁO CÁO DẠNG EXCEL
# ============================================================
@reports_bp.route("/export_report/excel")
@require_role("admin", "hr")
def export_report_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "Báo cáo tổng quan"

    ws.append(["Chỉ tiêu", "Giá trị"])
    ws.append(["Tổng nhân viên", 6])
    ws.append(["Tổng phòng ban", 4])
    ws.append(["Tỉ lệ chấm công", "96.5%"])
    ws.append(["Giờ TB/Tuần", 4])
    ws.append(["Tổng số ca", 1])
    ws.append(["Tổng lương", 0])

    # Auto-fit độ rộng
    for col in ws.columns:
        max_len = max(len(str(cell.value)) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max_len + 4

    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="BaoCaoTongQuan.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
